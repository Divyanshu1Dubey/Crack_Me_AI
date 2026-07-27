"""DOCX parser.

This is the workhorse for `cms_exclusive_material/`. Detects the format
inside the file and dispatches to the right extractor:

* `ClassicMCQExtractor`  — `Q1. ... / A. ... / B. ... / Answer: B / Explanation: ...`
* `BoxedMCQExtractor`    — Word tables shaped like `Question / Type / Option / Option / ... / Solution`
* `StatementMCQExtractor`— "I. ... II. ... III. ... / Select using the code below"
* `TheoryExtractor`      — long-form notes / headings / indexes

Format detection is heuristic (cheap) and runs before extraction. All
extractors are pure functions over a list of paragraphs / tables so they
can also be unit-tested directly.
"""
from __future__ import annotations

import io
import re
import time
import zipfile
from pathlib import Path
from typing import Iterable, List, Sequence, Tuple

from .dataclasses import ParsedDocument, ParsedImage, ParsedQuestion, ParsedTheory
from .docx_fidelity import (
    Paragraph,
    Table as FidelityTable,
    TableCell,
    TableRow,
    associate_images_to_questions,
    read_document_with_fidelity,
    render_blocks,
    render_paragraph,
    render_table,
)
from .text_utils import (
    clean_text,
    content_hash,
    extract_year_hint,
    flatten_paragraphs,
    is_likely_question_start,
    split_option_line,
)

# ----- Format detection ------------------------------------------------------

_QNUMBER_RE = re.compile(r"^Q\s*\d+[\.\):]", re.IGNORECASE)
_ANSWER_RE = re.compile(r"^\s*Answer\s*[:\-]\s*([A-Da-d])\b", re.IGNORECASE)
_EXPLANATION_RE = re.compile(r"^\s*Explanation\s*[:\-]\s*(.*)", re.IGNORECASE)
_SELECT_CODE_RE = re.compile(r"select using the code", re.IGNORECASE)
_ROMAN_HEAD_RE = re.compile(r"^\s*[IVX]{1,4}\.\s+[A-Z]", re.IGNORECASE)
_OPTION_LABEL_RE = re.compile(r"^\s*\(?([A-D])\)?[\.\):]\s+")


def _looks_classic_mcq(paragraphs: Sequence[str]) -> bool:
    scored = 0
    for p in paragraphs[:200]:
        if _QNUMBER_RE.match(p):
            scored += 2
        if _ANSWER_RE.match(p):
            scored += 2
        if _EXPLANATION_RE.match(p):
            scored += 1
        if _OPTION_LABEL_RE.match(p):
            scored += 1
    return scored >= 6


def _looks_boxed_mcq(paragraphs: Sequence[str], tables: Sequence[Sequence[Sequence[str]]]) -> bool:
    """Detect the Question/Type/Option/correct/incorrect pattern from Word tables.

    The flag column in Meduraa templates is column 2 (not column 0), so we
    have to look across all cells of every row to find the markers.
    """
    if not tables:
        return False
    for t in tables[:30]:
        if not t:
            continue
        col0 = [row[0].strip().lower().rstrip() if row else "" for row in t]
        if not col0 or col0[0] != "question":
            continue
        saw_option = False
        saw_correct = False
        for row in t:
            if not row:
                continue
            for cell in row:
                c = cell.strip().lower() if cell else ""
                if c == "option":
                    saw_option = True
                elif c in ("correct", "incorrect"):
                    saw_correct = True
        if saw_option and saw_correct:
            return True
    return False


def _looks_statement_mcq(paragraphs: Sequence[str]) -> bool:
    saw_roman = False
    saw_code = False
    for p in paragraphs[:60]:
        if _ROMAN_HEAD_RE.match(p):
            saw_roman = True
        if _SELECT_CODE_RE.search(p):
            saw_code = True
    return saw_roman and saw_code


def detect_format(
    paragraphs: Sequence[str],
    tables: Sequence[Sequence[Sequence[str]]],
) -> str:
    """Return one of the ParsedDocument.detected_type values."""
    if _looks_boxed_mcq(paragraphs, tables):
        return "mcq_boxed"
    if _looks_classic_mcq(paragraphs):
        return "mcq_classic"
    if _looks_statement_mcq(paragraphs):
        return "mcq_statement"
    # Theory: any non-trivial document with few/no MCQ markers.
    if not paragraphs:
        return "unknown"
    nonblank = sum(1 for p in paragraphs if p.strip())
    if nonblank >= 25:
        return "theory"
    return "unknown"


# ----- DOCX reader ----------------------------------------------------------

def _is_docx_namespace_error(exc: Exception) -> bool:
    """Detect failures rooted in undeclared xmlns / undeclared namespace prefix.

    Both lxml (which python-docx uses) and the stdlib xml parser raise
    errors when a document declares a prefix like ``w:nsid`` but never
    binds ``w`` to a namespace. We want to retry those with our own
    namespace-tolerant extractor instead of giving up on the file.
    """
    msg = str(exc).lower()
    needles = (
        "namespace prefix",
        "undeclared prefix",
        "undefined prefix",
        "is not declared",
        "namespacenotdeclared",
        "prefix '",
    )
    return any(n in msg for n in needles)


def _docx_read_xml_fallback(raw: bytes):
    """Last-resort extractor: pull paragraphs + tables straight from `word/document.xml`.

    Used only when ``python-docx`` blows up on namespace declarations (the
    underlying lxml invocation is too strict for documents like
    ``merged_notes-document (1).docx`` whose ``w14``/``w15`` prefixes go
    undeclared). The shape returned matches ``_docx_read`` so the rest of the
    pipeline sees no difference.
    """
    import xml.etree.ElementTree as ET

    # Only declare the namespace we actually need; we read by local-name to
    # tolerate any other (or missing) prefix binding.
    NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    paragraphs: List[str] = []
    tables: List[List[List[str]]] = []

    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as z:
            if "word/document.xml" not in z.namelist():
                return paragraphs, tables
            xml_bytes = z.read("word/document.xml")
    except Exception:
        return paragraphs, tables

    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError:
        return paragraphs, tables

    def _local(tag: str) -> str:
        return tag.split("}", 1)[-1] if "}" in tag else tag

    def _para_text(p_elem) -> str:
        parts: List[str] = []
        for t in p_elem.iter():
            local = _local(t.tag)
            if local == "t":
                if t.text:
                    parts.append(t.text)
            elif local == "tab":
                parts.append("	")
            elif local == "br":
                parts.append("\n")
        return "".join(parts)

    body = root.find("w:body", NS)
    if body is None:
        for child in root:
            if _local(child.tag) == "body":
                body = child
                break
    if body is None:
        return paragraphs, tables

    for child in body:
        local = _local(child.tag)
        if local == "p":
            paragraphs.append(_para_text(child))
        elif local == "tbl":
            rows: List[List[str]] = []
            for tr in child:
                if _local(tr.tag) != "tr":
                    continue
                cells: List[str] = []
                for tc in tr:
                    if _local(tc.tag) != "tc":
                        continue
                    cell_paras = [_para_text(p) for p in tc if _local(p.tag) == "p"]
                    cells.append("\n".join(cell_paras))
                rows.append(cells)
            if rows:
                tables.append(rows)
    return paragraphs, tables


def _docx_read(path: str):
    """Return (paragraphs, tables, raw_zip_bytes).

    Strategy:
      1. Try ``python-docx`` first (rich, but intolerant of unknown namespaces).
      2. If that explodes with a *namespace-prefix* error — which is the mode
         we have seen in the wild for ``merged_notes-document (1).docx`` and
         similar — retry with a hand-rolled XML extractor that only walks
         ``<w:p>`` / ``<w:tbl>`` and tolerates unknown namespace prefixes.
      3. Anything else propagates and turns into a ParsedDocument.errors entry.
    """
    with open(path, "rb") as fh:
        raw = fh.read()
    try:
        from docx import Document  # python-docx is heavy — import lazily
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(f"python-docx is required for DOCX parsing ({exc})") from exc

    try:
        doc = Document(io.BytesIO(raw))
        paragraphs = [p.text for p in doc.paragraphs]
        tables: List[List[List[str]]] = []
        for t in doc.tables:
            rows: List[List[str]] = []
            for row in t.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append(rows)
        return paragraphs, tables, raw, False
    except Exception as exc:
        if _is_docx_namespace_error(exc):
            paragraphs, tables = _docx_read_xml_fallback(raw)
            return paragraphs, tables, raw, True
        raise


def _extract_images_from_zip(raw: bytes) -> List[ParsedImage]:
    """Pull every embedded media file from the DOCX zip."""
    images: List[ParsedImage] = []
    with zipfile.ZipFile(io.BytesIO(raw)) as z:
        for name in z.namelist():
            if not name.startswith("word/media/"):
                continue
            data = z.read(name)
            ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
            mime = {
                "png": "image/png",
                "jpg": "image/jpeg",
                "jpeg": "image/jpeg",
                "gif": "image/gif",
                "webp": "image/webp",
                "svg": "image/svg+xml",
                "wmf": "image/x-wmf",
                "emf": "image/x-emf",
            }.get(ext, "application/octet-stream")
            w = h = 0
            if ext in ("png", "jpg", "jpeg", "gif", "webp"):
                try:
                    from PIL import Image
                    with Image.open(io.BytesIO(data)) as im:
                        w, h = im.size
                except Exception:
                    pass
            images.append(
                ParsedImage(
                    filename=Path(name).name,
                    raw_bytes=data,
                    mime_type=mime,
                    width=w,
                    height=h,
                    source_ref=name,
                )
            )
    return images


# ----- Extractors ------------------------------------------------------------

class ClassicMCQExtractor:
    """Parse `Q1. .../A./B./C./D./Answer: B/Explanation: ...` blocks."""

    def extract(self, paragraphs: Sequence[str], warnings_out: list | None = None) -> List[ParsedQuestion]:
        out: List[ParsedQuestion] = []
        warnings = warnings_out if warnings_out is not None else []
        # Collapse non-blank lines into a single token stream.
        cleaned = [clean_text(p) for p in paragraphs if p and p.strip()]
        if not cleaned:
            return out

        # Walk paragraphs, accumulate a current question.
        i = 0
        pos = 0
        while i < len(cleaned):
            line = cleaned[i]
            if not _QNUMBER_RE.match(line):
                i += 1
                continue
            # Initialize a new question.
            q = ParsedQuestion(
                position_index=pos,
                question_text="",
                raw_text=line,
            )
            # Consume the question stem (lines until we see an option label).
            stem = [line]
            i += 1
            while i < len(cleaned) and not _OPTION_LABEL_RE.match(cleaned[i]) and not _QNUMBER_RE.match(cleaned[i]):
                stem.append(cleaned[i])
                i += 1
            # The stem may include the label of the question in forms like "Q1. text A. opt"
            # Split off the prefix.
            full_stem = " ".join(stem).strip()
            q.question_text = re.sub(r"^Q\s*\d+[\.\):]\s*", "", full_stem, flags=re.IGNORECASE).strip()

            # Read options A/B/C/D (in any order; canonicalize).
            opts: dict[str, str] = {}
            while i < len(cleaned) and (_OPTION_LABEL_RE.match(cleaned[i]) or _QNUMBER_RE.match(cleaned[i])):
                if _QNUMBER_RE.match(cleaned[i]):
                    break
                m = _OPTION_LABEL_RE.match(cleaned[i])
                if not m:
                    break
                letter = m.group(1).upper()
                # Option text may span multiple lines until next option/answer/question.
                buf = [cleaned[i][m.end():].strip()]
                i += 1
                while i < len(cleaned) and not _OPTION_LABEL_RE.match(cleaned[i]) \
                        and not _ANSWER_RE.match(cleaned[i]) \
                        and not _EXPLANATION_RE.match(cleaned[i]) \
                        and not _QNUMBER_RE.match(cleaned[i]):
                    buf.append(cleaned[i])
                    i += 1
                opts[letter] = clean_text(" ".join(buf))
            q.option_a = opts.get("A", "")
            q.option_b = opts.get("B", "")
            q.option_c = opts.get("C", "")
            q.option_d = opts.get("D", "")

            # Answer: <letter>
            if i < len(cleaned) and _ANSWER_RE.match(cleaned[i]):
                m = _ANSWER_RE.match(cleaned[i])
                q.correct_answer = m.group(1).upper()
                i += 1
            # Explanation: <text>
            if i < len(cleaned) and _EXPLANATION_RE.match(cleaned[i]):
                m = _EXPLANATION_RE.match(cleaned[i])
                buf = [m.group(1).strip()]
                i += 1
                while i < len(cleaned) and not _QNUMBER_RE.match(cleaned[i]) and not _ANSWER_RE.match(cleaned[i]):
                    buf.append(cleaned[i])
                    i += 1
                q.explanation = clean_text(" ".join(buf))

            # W5 fix: require all 4 options + a correct marker before emitting.
            missing = [k for k in "ABCD" if not (opts.get(k, "") or "").strip()]
            if q.question_text and not missing and q.correct_answer in "ABCD":
                out.append(q)
                pos += 1
            elif q.question_text:
                reasons: list[str] = []
                if missing:
                    reasons.append(f"missing_options={','.join(missing)}")
                if q.correct_answer not in "ABCD":
                    reasons.append("missing_correct_marker")
                warnings.append(f"classic_question_skipped pos={pos}: {','.join(reasons)}")
        return out


class BoxedMCQExtractor:
    """Parse the `Question/Type/Option/correct/incorrect/Solution/Marks` table layout.

    The DOCX tables produced by Meduraa's authoring template have either 2
    or 3 columns depending on the row:

        Question | <question text>               (label/text)
        Type     | multiple_choice               (label/text)
        Option   | <option text>    | correct    (label/text/flag)
        Option   | <option text>    | incorrect
        Solution | <solution text>
        Marks    | 3                | 1          (label/text/weight?)

    This extractor normalizes that and produces a `ParsedQuestion` per table.
    """

    def extract(self, tables: Sequence[Sequence[Sequence[str]]], warnings_out: list | None = None) -> List[ParsedQuestion]:
        out: List[ParsedQuestion] = []
        pos = 0
        warnings = warnings_out if warnings_out is not None else []
        for t_idx, table in enumerate(tables):
            if not table or len(table) < 4:
                continue
            col0 = [row[0].strip().lower() if row else "" for row in table]
            if not col0 or col0[0] != "question":
                continue
            cells: dict[str, list[str]] = {}
            for row in table:
                if not row:
                    continue
                label = row[0].strip().lower()
                if not label:
                    continue
                # In a 3-column row the second cell is the value and the
                # third is the flag (correct/incorrect). In a 2-column row
                # the second cell IS the value; the third doesn't exist.
                text = row[1].strip() if len(row) >= 2 else ""
                if label == "option":
                    flag = row[2].strip().lower() if len(row) >= 3 else ""
                    cells.setdefault("option", []).append(text)
                    cells.setdefault("option_correct", []).append(flag)
                elif label == "marks":
                    cells.setdefault("marks", []).append(text)
                elif label.startswith("negative"):
                    cells.setdefault("negative", []).append(text)
                elif label == "solution":
                    cells.setdefault("solution", []).append(text)
                elif label == "question":
                    cells.setdefault("question", []).append(text)
                elif label.startswith("type"):
                    cells.setdefault("type", []).append(text)
            question_text = clean_text(" ".join(cells.get("question", [])))
            if not question_text:
                continue
            options = cells.get("option", [])
            correct_flags = cells.get("option_correct", [])
            if len(options) < 4:
                warnings.append(f"boxed_table={t_idx} fewer_than_4_options: found {len(options)}")
                continue
            if len(options) > 4:
                warnings.append(f"boxed_table={t_idx} more_than_4_options_truncated: found {len(options)}")
            opts = options[:4]
            # Stricter flag check (W4 fix): exact match only.
            correct = ""
            unknown_flags: list[int] = []
            for idx, flag in enumerate(correct_flags[:4]):
                if flag == "correct":
                    correct = "ABCD"[idx]
                elif flag and flag != "incorrect":
                    unknown_flags.append(idx)
            if not correct:
                warnings.append(f"boxed_table={t_idx} missing_correct_marker")
            if unknown_flags:
                warnings.append(f"boxed_table={t_idx} unknown_option_flag={unknown_flags}")
            q = ParsedQuestion(
                position_index=pos,
                question_text=question_text,
                option_a=clean_text(opts[0]),
                option_b=clean_text(opts[1]),
                option_c=clean_text(opts[2]),
                option_d=clean_text(opts[3]),
                correct_answer=correct,
                explanation=clean_text(" ".join(cells.get("solution", []))),
                raw_text="\n".join(", ".join(r) for r in table),
                extra={"marks": (cells.get("marks") or ["1"])[0].strip() or "1"},
            )
            try:
                q.marks = int(q.extra["marks"].split()[0])
            except (ValueError, IndexError):
                q.marks = 1
            # Parse negative_marks from the "Negative" row (e.g. "1" / "0.33" / "-0.33").
            neg_raw = (cells.get("negative") or ["0"])[0].strip()
            try:
                q.negative_marks = abs(float(re.sub(r"[^0-9.\-]", "", neg_raw) or "0"))
            except (ValueError, IndexError):
                q.negative_marks = 0.0
            out.append(q)
            pos += 1
        return out


class BoxedMCQExtractorFidelity:
    """High-fidelity boxed extractor.

    Consumes a list of fidelity-aware ``docx_fidelity.Table`` objects plus a
    flat list of paragraphs. Reuses the string-based extractor to recover
    the structural fields (question/option A-D/correct/explanation/marks)
    and then **upgrades** each question with HTML-rendered versions of those
    fields so colours, bold, italic, highlights, lists, and tables round-trip.

    Per-question image association is rebuilt from the fidelity walk —
    only images that sit inside the boxed table belonging to a question
    are recorded in that question's ``image_refs``.
    """

    @staticmethod
    def _row_label(row: TableRow) -> str:
        """Return the lowercased label of the first cell in the row (e.g. 'option')."""
        if not row.cells:
            return ""
        first = row.cells[0]
        return "".join(r.text for p in first.paragraphs for r in p.runs).strip().lower().rstrip()

    @staticmethod
    def _row_value_html(row: TableRow, image_url_for: dict | None) -> str:
        """Render the *value* cells (everything except column 0) as HTML."""
        if not row.cells:
            return ""
        parts: List[str] = []
        for idx, cell in enumerate(row.cells):
            if idx == 0:
                continue
            parts.append(render_blocks([__import__("material_importer.parser.docx_fidelity", fromlist=["DocumentBlock"]).DocumentBlock(paragraph=p) for p in cell.paragraphs], image_url_for=image_url_for) if False else render_blocks(
                [
                    __import__("material_importer.parser.docx_fidelity", fromlist=["DocumentBlock"]).DocumentBlock(paragraph=p)
                    for p in cell.paragraphs
                ],
                image_url_for=image_url_for,
            ))
        return "\n".join(p for p in parts if p.strip())

    def extract(
        self,
        tables: Sequence[FidelityTable],
        warnings_out: list | None = None,
        image_url_for: dict | None = None,
        block_paragraph_indexes: List[int] | None = None,
    ) -> List[ParsedQuestion]:
        from .docx_fidelity import DocumentBlock  # local import to avoid cycle
        warnings = warnings_out if warnings_out is not None else []

        # First pass: convert each fidelity Table to a string table (legacy format)
        # so the existing BoxedMCQExtractor can extract structural fields.
        legacy_tables: List[List[List[str]]] = []
        for t in tables:
            legacy_rows: List[List[str]] = []
            for row in t.rows:
                legacy_row: List[str] = []
                for cell in row.cells:
                    legacy_row.append("".join(r.text for p in cell.paragraphs for r in p.runs).strip())
                legacy_rows.append(legacy_row)
            legacy_tables.append(legacy_rows)

        legacy = BoxedMCQExtractor().extract(legacy_tables, warnings_out=warnings)
        # Map: each legacy ParsedQuestion was built from `tables[i]` where i == pos.
        # Re-render each field using the fidelity tables.
        from .docx_fidelity import render_blocks as _render_blocks  # already imported

        for pos, q in enumerate(legacy):
            if pos >= len(tables):
                break
            t = tables[pos]
            if block_paragraph_indexes and pos < len(block_paragraph_indexes):
                q.paragraph_index = block_paragraph_indexes[pos]
            # Per-question image set: collect images that sit inside this table.
            image_filenames: List[str] = []
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            if r.image_filename and r.image_filename not in image_filenames:
                                image_filenames.append(r.image_filename)
            q.extra["images"] = image_filenames
            q.image_refs = image_filenames
            # Render each row's HTML to upgrade the textual fields.
            html_question: List[str] = []
            html_option = {"A": [], "B": [], "C": [], "D": []}
            html_solution: List[str] = []
            for row in t.rows:
                label = self._row_label(row)
                value_html = self._row_value_html(row, image_url_for)
                if label == "question":
                    html_question.append(value_html)
                elif label.startswith("option") and len(row.cells) >= 2:
                    # Map row to A/B/C/D by order of appearance.
                    option_letter = ["A", "B", "C", "D"][len([k for k in html_option if html_option[k]]) % 4]
                    html_option[option_letter].append(value_html)
                elif label == "solution":
                    html_solution.append(value_html)
            if html_question:
                q.question_text = "\n".join(html_question) or q.question_text
            for letter, htmls in html_option.items():
                if htmls:
                    setattr(q, f"option_{letter.lower()}", "\n".join(htmls) or getattr(q, f"option_{letter.lower()}"))
            if html_solution:
                q.explanation = "\n".join(html_solution) or q.explanation
        return legacy


class StatementMCQExtractor:
    """Parse statement-based MCQ (I/II/III/IV statements + 'Select using the code below')."""

    _STATEMENT_RE = re.compile(r"^[IVX]{1,4}[\.\)]\s+", re.IGNORECASE)
    _CODE_OPTION_RE = re.compile(r"^\s*\(?([A-D])\)?[\.\):\-]?\s*(.*)$")
    _ANSWER_INLINE_RE = re.compile(r"\bCorrect\s*[:\-]\s*([A-D])\b", re.IGNORECASE)

    def extract(self, paragraphs: Sequence[str]) -> List[ParsedQuestion]:
        out: List[ParsedQuestion] = []
        cleaned = [clean_text(p) for p in paragraphs if p and p.strip()]
        # Find question starts: a paragraph that isn't a statement, isn't an option code,
        # and is followed by I./II. style statements.
        i = 0
        pos = 0
        while i < len(cleaned):
            # Look for pattern: stem (1+ lines) -> statements -> code prompt -> options
            win = cleaned[i:i + 8]
            if not win:
                break
            # Heuristic: skip if next 1-4 lines are I/II/III numbered statements.
            stmt_starts = sum(1 for w in win if self._STATEMENT_RE.match(w))
            if stmt_starts < 2:
                i += 1
                continue
            # Find the statement block range.
            stem_lines = []
            while i < len(cleaned) and not self._STATEMENT_RE.match(cleaned[i]):
                stem_lines.append(cleaned[i])
                i += 1
            question_text = clean_text(" ".join(stem_lines))
            statements = []
            while i < len(cleaned) and self._STATEMENT_RE.match(cleaned[i]):
                statements.append(cleaned[i])
                i += 1
            # Skip "Select using the code below" prompt.
            if i < len(cleaned) and _SELECT_CODE_RE.search(cleaned[i]):
                i += 1
            # Read up to 4 options.
            opts: dict[str, str] = {}
            while i < len(cleaned) and len(opts) < 4:
                m = self._CODE_OPTION_RE.match(cleaned[i])
                if not m:
                    break
                letter = m.group(1).upper()
                if letter in "ABCD":
                    opts[letter] = clean_text(m.group(2))
                i += 1
            # Optional inline correct marker.
            correct = ""
            for k in ("A", "B", "C", "D"):
                for line in cleaned[max(0, i - 4): i + 1]:
                    if f"Correct: {k}" in line or f"Correct-{k}" in line:
                        correct = k
                        break
                if correct:
                    break
            if question_text and len(opts) >= 2:
                q = ParsedQuestion(
                    position_index=pos,
                    question_text=question_text,
                    option_a=opts.get("A", ""),
                    option_b=opts.get("B", ""),
                    option_c=opts.get("C", ""),
                    option_d=opts.get("D", ""),
                    correct_answer=correct,
                    raw_text="\n".join(stem_lines + statements),
                    extra={"statements": statements},
                )
                out.append(q)
                pos += 1
        return out


class TheoryExtractor:
    """Extract headings, paragraphs, lists, and tables from theory notes."""

    _HEADING_KEYWORDS = (
        "NOTES", "INDEX", "CHAPTER", "CH ", "UNIT ", "TOPIC",
        "BY DR", "DR.",
    )
    _LIST_RE = re.compile(r"^\s*[•◦‣\-\*]\s+", re.IGNORECASE)
    _NUM_LIST_RE = re.compile(r"^\s*\d+[\.\)]\s+")

    def extract(self, paragraphs: Sequence[str], tables: Sequence[Sequence[Sequence[str]]]) -> List[ParsedTheory]:
        out: List[ParsedTheory] = []
        cleaned = flatten_paragraphs(paragraphs)
        pos = 0
        current_heading = ""
        current_subheading = ""
        for p in cleaned:
            upper = p.upper()
            is_heading = any(upper.startswith(k) or upper == k.strip() for k in self._HEADING_KEYWORDS) or p.isupper()
            long_topic = p.isupper() and 5 < len(p) < 80
            if is_heading and len(p) < 80:
                current_heading = p.strip()
                continue
            if long_topic:
                current_subheading = p.strip()
                continue
            block_type = "paragraph"
            if self._LIST_RE.match(p) or self._NUM_LIST_RE.match(p):
                block_type = "list"
            elif upper == "INDEX" or upper.startswith("INDEX "):
                block_type = "index"
            theory = ParsedTheory(
                position_index=pos,
                heading=current_heading,
                subheading=current_subheading,
                body_text=p,
                block_type=block_type,
                raw_text=p,
                keywords=self._extract_keywords(p),
            )
            out.append(theory)
            pos += 1
        # Convert tables to typed blocks.
        for t_idx, table in enumerate(tables):
            if not table:
                continue
            joined = "\n".join(" | ".join(c.strip() for c in row) for row in table if row)
            if not joined.strip():
                continue
            out.append(
                ParsedTheory(
                    position_index=pos,
                    heading=current_heading,
                    subheading=current_subheading,
                    body_text=joined,
                    block_type="table",
                    raw_text=joined,
                    extra={"table_index": t_idx, "rows": len(table)},
                )
            )
            pos += 1
        return out

    @staticmethod
    def _extract_keywords(text: str) -> List[str]:
        if not text:
            return []
        # Cheap keyword extractor — top longish words, stopwords removed.
        STOP = {
            "the", "and", "for", "with", "from", "this", "that", "are", "was",
            "were", "have", "has", "had", "but", "not", "you", "your", "can",
            "may", "see", "use", "used", "using", "such", "any", "all", "most",
            "more", "less", "than", "into", "out", "over", "under", "between",
            "also", "its", "their", "his", "her", "she", "they", "them", "our",
            "we", "be", "is", "of", "in", "on", "to", "by", "an", "a", "as",
            "at", "it", "or", "if", "no", "do", "so", "up", "very",
        }
        words = re.findall(r"[A-Za-z][A-Za-z\-]{3,}", text)
        seen = set()
        out: List[str] = []
        for w in words:
            lw = w.lower()
            if lw in STOP or lw in seen:
                continue
            seen.add(lw)
            out.append(w)
            if len(out) >= 8:
                break
        return out


# ----- Public entry point ----------------------------------------------------

class DOCXParser:
    """Orchestrator: read DOCX → detect format → dispatch → return ParsedDocument."""

    def __init__(self) -> None:
        self.classic = ClassicMCQExtractor()
        self.boxed = BoxedMCQExtractor()
        self.boxed_fidelity = BoxedMCQExtractorFidelity()
        self.statement = StatementMCQExtractor()
        self.theory = TheoryExtractor()

    def parse(self, path: str) -> ParsedDocument:
        start = time.time()
        fn = Path(path).name
        try:
            paragraphs, tables, raw, used_xml_fallback = _docx_read(path)
        except Exception as exc:
            return ParsedDocument(
                source_filename=fn,
                file_format="docx",
                detected_type="unknown",
                parser_used="docx",
                errors=[f"Failed to open DOCX: {exc}"],
            )
        images = _extract_images_from_zip(raw)
        detected = detect_format(paragraphs, tables)
        doc = ParsedDocument(
            source_filename=fn,
            file_format="docx",
            detected_type=detected,
            parser_used="docx",
            images=images,
        )
        # Build fidelity blocks (paragraphs + tables with formatting preserved).
        try:
            fidelity_blocks = read_document_with_fidelity(path)
        except Exception as exc:
            fidelity_blocks = []
            doc.warnings.append(f"fidelity_read_failed: {exc}")
        fidelity_tables: List[FidelityTable] = [b.table for b in fidelity_blocks if b.table is not None]
        fidelity_table_block_indexes: List[int] = [b.paragraph_index for b in fidelity_blocks if b.table is not None]
        doc.meta["fidelity_block_count"] = len(fidelity_blocks)
        if used_xml_fallback:
            doc.warnings.append("xml_fallback_used: namespace-tolerant XML extractor used (python-docx could not parse)")
            doc.parser_used = "docx.xml_fallback"
        if detected == "mcq_boxed":
            doc.questions = self.boxed_fidelity.extract(fidelity_tables, warnings_out=doc.warnings, block_paragraph_indexes=fidelity_table_block_indexes)
            # If fidelity returned nothing, fall back to the string-based extractor.
            if not doc.questions:
                doc.questions = self.boxed.extract(tables, warnings_out=doc.warnings)
            for q in doc.questions:
                q.extra.setdefault("parser_schema", "boxed")
            # Some boxed files also embed free text + tables. If <2 questions
            # were recovered, fall back to classic extraction.
            if len(doc.questions) < 2:
                extra = self.classic.extract(paragraphs, warnings_out=doc.warnings)
                for eq in extra:
                    eq.extra.setdefault("parser_schema", "classic")
                if extra:
                    doc.questions.extend(extra)
            doc.parser_used = ("docx.xml_fallback+" if used_xml_fallback else "docx.") + "boxed+classic"
        elif detected == "mcq_classic":
            doc.questions = self.classic.extract(paragraphs, warnings_out=doc.warnings)
            for q in doc.questions:
                q.extra.setdefault("parser_schema", "classic")
            doc.parser_used = ("docx.xml_fallback+" if used_xml_fallback else "docx.") + "classic"
        elif detected == "mcq_statement":
            doc.questions = self.statement.extract(paragraphs)
            for q in doc.questions:
                q.extra.setdefault("parser_schema", "statement")
            doc.parser_used = ("docx.xml_fallback+" if used_xml_fallback else "docx.") + "statement"
        elif detected == "theory":
            doc.theory_blocks = self.theory.extract(paragraphs, tables)
            doc.parser_used = ("docx.xml_fallback+" if used_xml_fallback else "docx.") + "theory"
        else:
            # Unknown but try anyway — some unknown files still contain MCQs.
            doc.questions = self.classic.extract(paragraphs, warnings_out=doc.warnings)
            for q in doc.questions:
                q.extra.setdefault("parser_schema", "classic")
            doc.theory_blocks = self.theory.extract(paragraphs, tables)
            doc.parser_used = ("docx.xml_fallback+" if used_xml_fallback else "docx.") + "heuristic"
            if doc.questions:
                doc.detected_type = "mcq_classic"
            elif doc.theory_blocks:
                doc.detected_type = "theory"
            else:
                doc.warnings.append("Could not classify content; no MCQ or theory recovered.")

        # Files that mix headings + indexes may have a leading INDEX block. Strip
        # obvious front-matter from the first theory block if it is the entire
        # body — keeps the INDEX out of the long-form theory corpus.
        if doc.theory_blocks and doc.theory_blocks[0].body_text.upper().strip() == "INDEX":
            doc.theory_blocks.pop(0)
            for idx, t in enumerate(doc.theory_blocks):
                t.position_index = idx

        # Year hint from any question text.
        all_text = "\n".join((q.question_text or "") for q in doc.questions)
        if all_text:
            doc.meta["year"] = extract_year_hint(all_text)

        # Re-number position_index so it's dense.
        for idx, q in enumerate(doc.questions):
            q.position_index = idx
        for idx, t in enumerate(doc.theory_blocks):
            t.position_index = idx

        doc.duration_ms = int((time.time() - start) * 1000)
        return doc


# ----- Top-level alias -------------------------------------------------------

def parse_docx(path: str) -> ParsedDocument:
    return DOCXParser().parse(path)
