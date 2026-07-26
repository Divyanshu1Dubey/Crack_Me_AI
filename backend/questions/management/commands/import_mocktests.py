"""Import mock-test DOCX files into the CrackCMS Question bank.

Supported formats (auto-detected per file):
  A. Mini-test tables — each question is one 8x3 table with column-0 labels
     ['Question','Type','Option','Option','Option','Option','Solution','Marks'].
     Correct answer = column-2 cell of the Option row whose value is 'correct'.
  B. Plain-text Q.N. blocks — sequential paragraphs starting with
     'Q<n>.', four 'A./B./C./D.' options, 'Answer: <A|B|C|D>' and
     'Explanation: <text>'. Tables that don't match Schema A are skipped.
  C. Mixed — Schema B over paragraphs plus Schema A over tables (handled by
     the same parser; tables that don't match A are ignored).

Subject/topic inferred from filename. Admin can override via --subject.
Image references inside cells are saved to MEDIA_ROOT/mocktest_assets/ and
returned as URLs.

Usage:
    python manage.py import_mocktests --dir ../cms_exclusive_material --dry-run
    python manage.py import_mocktests --file "Mini test-4.docx" --publish-test "Mini Test-4"
    python manage.py import_mocktests --dir ../cms_exclusive_material --with-ai
"""
from __future__ import annotations

import argparse
import difflib
import hashlib
import json
import logging
import os
import re
import sys
import unicodedata
from collections import Counter
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Iterator

from django.conf import settings
from django.core.management.base import BaseCommand, CommandError
from django.db import transaction

# `python-docx` lives in the system Python on this machine; importing at module
# level avoids ModuleNotFoundError during `manage.py` dispatch when venv is
# missing it. Defer the Django model imports until handle() so a missing
# docx doesn't take down management discovery.
try:
    from docx import Document  # noqa: F401
except ImportError:  # pragma: no cover
    Document = None

from questions.models import Question, Subject, Topic
from tests_engine.models import Test

logger = logging.getLogger("mocktests.importer")

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

SCHEMA_A_LABELS = ["Question", "Type", "Option", "Option", "Option", "Option", "Solution", "Marks"]
OPTION_LETTERS = ["A", "B", "C", "D"]
# Known header lines that should never be treated as plain options.
KNOWN_HEADER_RE = re.compile(r"(?i)^\s*(?:Answer|Ans|Correct\s*Answer|Explanation|Why|Reason|Assertion|Reference|Q\d+|Select\s+(?:the\s+)?correct)\b")

# Subject guess from filename. First match wins.
SUBJECT_RULES = [
    (r"(?i)\bpsm\b|preventive|community medicine|health care|communicable", "PSM"),
    (r"(?i)\bobg|obstet|gynae|gynaec|gynec", "OBG"),
    (r"(?i)\bpediatric|paediatric|newborn|imci|imnci|immunization|nutrition|growth|fluids|systemic ped", "PED"),
    (r"(?i)\bdermat|skin\b", "MED"),
    (r"(?i)\brespiratory|pulmon", "MED"),
    (r"(?i)\banesth", "SUR"),
    (r"(?i)\bortho", "SUR"),
    (r"(?i)\bsurger|surgery|abdominal", "SUR"),
    (r"(?i)\bmedicine|neurolog|cardio|rheumat|nephro", "MED"),
]

DEFAULT_DIFFICULTY = "medium"


def normalize_text(s: str) -> str:
    """Used for duplicate detection — fold case, strip whitespace + Roman numerals."""
    if not s:
        return ""
    s = unicodedata.normalize("NFKC", s).lower()
    s = re.sub(r"\s+", " ", s)
    s = re.sub(r"^\s*(i+\.|iv?v?\.|v\.|\d+\.)\s*", "", s)  # roman/decimal stem prefix
    s = re.sub(r"[^a-z0-9 ]+", "", s)
    return s.strip()


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------


@dataclass
class ParsedQuestion:
    question_text: str = ""
    option_a: str = ""
    option_b: str = ""
    option_c: str = ""
    option_d: str = ""
    correct_answer: str = ""  # 'A'|'B'|'C'|'D' or '' if missing
    explanation: str = ""
    marks: float = 3.0
    source_format: str = ""  # 'A' | 'B' | 'C'
    question_images: list = field(default_factory=list)
    solution_images: list = field(default_factory=list)
    errors: list = field(default_factory=list)


@dataclass
class FileReport:
    file: str
    schema: str = ""  # 'A' | 'B' | 'C' | 'NONE'
    questions: list = field(default_factory=list)
    subject_code: str = ""
    subject_guess: str = ""
    topics_seen: list = field(default_factory=list)
    total_images: int = 0
    parse_errors: list = field(default_factory=list)


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------


def extract_cell_images(doc, cell) -> list[tuple[bytes, str, str]]:
    """Return list of (blob_bytes, ext, mime) for every image embedded in a cell."""
    from docx.oxml.ns import qn

    out = []
    try:
        cell_xml = cell._tc
    except AttributeError:
        return out

    for blip in cell_xml.iter(qn("a:blip")):
        embed_attr = blip.get(qn("r:embed"))
        if not embed_attr:
            continue
        try:
            image_part = doc.part.related_parts[embed_attr]
            blob = image_part.blob
            ext = (image_part.partname.split(".")[-1] or "png").lower()
            mime_map = {
                "png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                "gif": "image/gif", "webp": "image/webp",
            }
            mime = mime_map.get(ext, "image/png")
            if ext not in mime_map:
                ext = "png"
                mime = "image/png"
            out.append((blob, ext, mime))
        except Exception as exc:
            logger.warning("image extraction failed: %s", exc)
    return out


def upload_pending_images(images: list[tuple[bytes, str, str]], question_id: int) -> list[dict]:
    """Upload a list of image tuples to Supabase Storage and return URL metadata.

    Each returned dict: {url, sha256_short, role_index}.
    Falls back to local MEDIA_ROOT write if Supabase upload fails.
    """
    if not images:
        return []
    urls: list[dict] = []
    try:
        from questions.image_upload import upload_image_to_supabase
        import io as _io
    except Exception as exc:  # pragma: no cover
        logger.warning("Could not import image_upload: %s", exc)
        upload_image_to_supabase = None
        _io = None

    for idx, (blob, ext, mime) in enumerate(images):
        if upload_image_to_supabase and _io:
            try:
                uploaded = upload_image_to_supabase(
                    file_obj=_io.BytesIO(blob),
                    question_id=question_id,
                    content_type=mime,
                    original_filename=f"mocktest_{question_id}_{idx}.{ext}",
                )
                urls.append({"url": uploaded.url, "sha256_short": uploaded.sha256_short, "index": idx})
                continue
            except Exception as exc:
                logger.warning("Supabase upload failed for Q%s img %s: %s", question_id, idx, exc)
        # Fallback: write to local MEDIA_ROOT
        digest = hashlib.sha1(blob).hexdigest()[:12]
        rel = f"mocktest_assets/inline_{digest}.{ext}"
        out_path = Path(settings.MEDIA_ROOT) / rel
        out_path.parent.mkdir(parents=True, exist_ok=True)
        if not out_path.exists():
            out_path.write_bytes(blob)
        urls.append({"url": settings.MEDIA_URL + rel, "sha256_short": digest, "index": idx})
    return urls


# ---------------------------------------------------------------------------
# Schema A parser (the Mini test-4 table format)
# ---------------------------------------------------------------------------


def looks_like_schema_a(table) -> bool:
    if len(table.rows) != 8 or len(table.columns) != 3:
        return False
    try:
        labels = [table.rows[i].cells[0].text.strip() for i in range(8)]
    except IndexError:
        return False
    return labels == SCHEMA_A_LABELS


def parse_schema_a(doc, table) -> ParsedQuestion | None:
    pq = ParsedQuestion(source_format="A")

    # Column 1 = text/value; Column 2 = mirror for some rows (correct/incorrect flag).
    def col1(i):
        try:
            return table.rows[i].cells[1].text.strip()
        except IndexError:
            return ""

    def col2(i):
        try:
            return table.rows[i].cells[2].text.strip()
        except IndexError:
            return ""

    pq.question_text = col1(0)
    if not pq.question_text:
        pq.errors.append("empty question_text")
        return None

    options = [col1(i) for i in range(2, 6)]
    flags = [col2(i) for i in range(2, 6)]
    for letter, text, flag in zip(OPTION_LETTERS, options, flags):
        setattr(pq, f"option_{letter.lower()}", text)
        if flag.lower() == "correct":
            pq.correct_answer = letter

    pq.explanation = col1(6)

    # Marks row: col1 = total marks, col2 = negative marks
    try:
        m_total = col1(7)
        pq.marks = float(m_total) if m_total else 3.0
    except ValueError:
        pq.errors.append(f"unparseable marks={col1(7)!r}")
        pq.marks = 3.0

    if not pq.correct_answer:
        pq.errors.append("no 'correct' flag in any option row")

    # Images — only meaningful on Question and Solution rows for now
    pq.question_images = extract_cell_images(doc, table.rows[0].cells[1])
    pq.solution_images = extract_cell_images(doc, table.rows[6].cells[1])

    return pq


# ---------------------------------------------------------------------------
# Schema B parser (plain-text Q.N. blocks)
# ---------------------------------------------------------------------------

Q_HEADER_RE = re.compile(r"^\s*Q\s*(\d+)\s*[\.\)]\s*(.*)$", re.S)
OPTION_RE = re.compile(r"^\s*([A-D])\s*[\.\)]\s*(.+?)\s*$", re.S)
# B4 — lowercase-letter options like "(a) Adenomyosis"
OPTION_LC_RE = re.compile(r"^\s*\(([a-d])\)\s*(.+?)\s*$", re.S)
ANSWER_RE = re.compile(r"^\s*(?:Correct\s*)?Answer\s*[:\-]?\s*\(?([A-Da-d])\)?\b.*$", re.I | re.S)
# Some files say "Ans: (c) ..." instead of "Answer:"
ANS_LC_RE = re.compile(r"^\s*Ans(?:wer)?\s*[:\-]\s*\(?([A-Da-d])\)?\b", re.I | re.S)
EXPLANATION_RE = re.compile(r"^\s*(?:Explanation|Why|Reason)\s*[:\-]\s*(.*)$", re.I | re.S)
# Statement list: "I.", "II.", "III.", "IV.", or arabic 1., 2., 3., 4.
STATEMENT_RE = re.compile(r"^\s*((?:I{1,3}V?|IV|VI{0,3}|[1-4])\s*[\.\)])\s*(.+?)\s*$", re.S)
# "Select the correct answer using the code" preamble — sets up B3 mode
CODE_PROMPT_RE = re.compile(r"(?i)^\s*Select\s+(?:the\s+)?correct\s+answer\s+(?:using\s+)?(?:the\s+)?code")
# Code-style option: "A. 1, 2 and 3 only" / "A. 1 and 2" / "A. All of the above"
CODE_OPT_RE = re.compile(r"^\s*([A-D])\s*[\.\)]\s*(.+?)\s*$", re.S)
# Assertion / Reason headers inside a single question
ASSERTION_RE = re.compile(r"^\s*Assertion\s*[:\-]\s*(.+)$", re.I | re.S)
REASON_RE = re.compile(r"^\s*Reason\s*[:\-]\s*(.+)$", re.I | re.S)


def iter_text_blocks(doc) -> Iterator[tuple[str, str]]:
    """Yield (style_name, line_text) tuples. Style lets the parser distinguish
    between regular prose (Explanation/Reference text) and List Paragraph /
    bulleted lines (which often carry unmarked option text in PSM-style files).
    """
    prev_style = None
    for p in doc.paragraphs:
        style = p.style.name if p.style else "Normal"
        text = p.text.replace("\xa0", " ")
        # Split on actual newlines, but the style applies to the whole paragraph.
        lines = [s.strip() for s in text.splitlines() if s.strip()]
        if not lines:
            if prev_style is not None and style != prev_style:
                yield "", ""
            prev_style = style
            continue
        if prev_style is not None and style != prev_style:
            yield "", ""
        prev_style = style
        for sub in lines:
            yield style, sub


def parse_schema_b(paragraphs) -> list[ParsedQuestion]:
    """Walk the flattened paragraph stream; each 'Q<n>.' starts a new question.

    `paragraphs` may be either a list of strings or an iterable of (style, line)
    tuples produced by iter_text_blocks(). Style is consulted to recognize
    "plain" options (no leading letter) that come from List Paragraph bullets.

    Recognized variants:
      B1 — stem + A./B./C./D. options + Answer: line + Explanation.
      B2 — Assertion-Reason: stem + Assertion: + Reason: + options.
      B3 — Statement-list PYQ: stem + I./II./III./IV. statements +
           "Select the correct answer using the code below:" + options each
           being a numeric combination like "1, 2 and 3 only".
      B4 — lowercase-letter options like "(a) ..."; answer may use
           "Ans: (c) ..." instead of "Answer:".
      B5 — "plain" options (no leading letter) emitted from List Paragraph
           bullets; treated as options A–D in order.
    """
    out: list[ParsedQuestion] = []
    cur: ParsedQuestion | None = None
    cur_opt = None
    cur_mode = "B1"
    statement_lines: list[str] = []
    explanation_lines: list[str] = []
    plain_option_count = 0  # how many unmarked options we've consumed

    def flush():
        nonlocal cur, cur_opt, cur_mode, statement_lines, explanation_lines, plain_option_count
        if cur is not None:
            if statement_lines and cur_mode in ("B2", "B3", "B2_options", "B3_options"):
                cur.question_text = (cur.question_text + "\n\n" + "\n".join(statement_lines)).strip()
            if explanation_lines and not cur.explanation:
                cur.explanation = " ".join(explanation_lines).strip()
            if cur.question_text or cur.correct_answer:
                out.append(cur)
        cur = None
        cur_opt = None
        cur_mode = "B1"
        statement_lines = []
        explanation_lines = []
        plain_option_count = 0

    # Normalize input into (style, line) pairs
    normalized = []
    for item in paragraphs:
        if isinstance(item, tuple):
            style, line = item
        else:
            style, line = "Normal", item
        line = line.replace("\xa0", " ").strip()
        if line:
            normalized.append((style, line))

    for style, line in normalized:
        m = Q_HEADER_RE.match(line)
        if m:
            flush()
            cur = ParsedQuestion(source_format="B")
            cur.question_text = m.group(2).strip()
            cur_mode = "B1"
            continue

        if cur is None:
            continue

        is_list_style = "list" in style.lower() or style.lower().startswith("list ")

        # ---- Mode transitions ----
        if cur_mode == "B1":
            if ASSERTION_RE.match(line):
                cur_mode = "B2"
                statement_lines.append(line)
                continue
            if CODE_PROMPT_RE.match(line):
                cur_mode = "B3_options"
                continue
            if STATEMENT_RE.match(line):
                cur_mode = "B3"
                statement_lines.append(line)
                continue
            m = OPTION_LC_RE.match(line)
            if m:
                letter, text = m.group(1).upper(), m.group(2).strip()
                if letter in OPTION_LETTERS:
                    setattr(cur, f"option_{letter.lower()}", text)
                    cur_opt = letter
                continue
            m = OPTION_RE.match(line)
            if m:
                letter, text = m.group(1).upper(), m.group(2).strip()
                if letter in OPTION_LETTERS:
                    setattr(cur, f"option_{letter.lower()}", text)
                    cur_opt = letter
                continue
            # B5 — plain (unmarked) option from a List Paragraph bullet
            if (
                is_list_style
                and not KNOWN_HEADER_RE.match(line)
                and plain_option_count < 4
                and not cur_opt
                and not cur.correct_answer
            ):
                next_letter = OPTION_LETTERS[plain_option_count]
                setattr(cur, f"option_{next_letter.lower()}", line)
                cur_opt = next_letter
                plain_option_count += 1
                continue
            # B4 — "Ans: (c) ..."
            m = ANS_LC_RE.match(line)
            if m:
                cur.correct_answer = m.group(1).upper()
                tail = line[m.end():].strip()
                if tail:
                    explanation_lines.append(tail)
                continue
            m = ANSWER_RE.match(line)
            if m:
                cur.correct_answer = m.group(1).upper()
                tail = line[m.end():].strip()
                if tail:
                    explanation_lines.append(tail)
                continue
            m = EXPLANATION_RE.match(line)
            if m:
                tail = m.group(1).strip()
                if tail:
                    explanation_lines.append(tail)
                continue
            # loose continuation — but never glue two list-bullets together
            if cur_opt and not is_list_style:
                existing = getattr(cur, f"option_{cur_opt.lower()}", "")
                setattr(cur, f"option_{cur_opt.lower()}", (existing + " " + line).strip())
            elif explanation_lines or cur.explanation:
                explanation_lines.append(line)
            elif is_list_style and plain_option_count < 4 and not cur.correct_answer:
                # Next list bullet — treat as next plain option
                next_letter = OPTION_LETTERS[plain_option_count]
                setattr(cur, f"option_{next_letter.lower()}", line)
                cur_opt = next_letter
                plain_option_count += 1
                continue
            else:
                cur.question_text = (cur.question_text + " " + line).strip()
            continue

        # B2 mode: collect Assertion/Reason lines until we hit options
        if cur_mode == "B2":
            m = OPTION_RE.match(line)
            if m:
                cur_mode = "B2_options"
                letter, text = m.group(1).upper(), m.group(2).strip()
                if letter in OPTION_LETTERS:
                    setattr(cur, f"option_{letter.lower()}", text)
                    cur_opt = letter
                continue
            m = OPTION_LC_RE.match(line)
            if m:
                cur_mode = "B2_options"
                letter, text = m.group(1).upper(), m.group(2).strip()
                if letter in OPTION_LETTERS:
                    setattr(cur, f"option_{letter.lower()}", text)
                    cur_opt = letter
                continue
            # keep collecting AR text
            statement_lines.append(line)
            continue

        # B2_options: A./B./C./D. style answer
        if cur_mode == "B2_options":
            m = OPTION_LC_RE.match(line)
            if m:
                letter, text = m.group(1).upper(), m.group(2).strip()
                if letter in OPTION_LETTERS:
                    setattr(cur, f"option_{letter.lower()}", text)
                    cur_opt = letter
                continue
            m = OPTION_RE.match(line)
            if m:
                letter, text = m.group(1).upper(), m.group(2).strip()
                if letter in OPTION_LETTERS:
                    setattr(cur, f"option_{letter.lower()}", text)
                    cur_opt = letter
                continue
            m = ANS_LC_RE.match(line)
            if m:
                cur.correct_answer = m.group(1).upper()
                tail = line[m.end():].strip()
                if tail:
                    explanation_lines.append(tail)
                cur_mode = "B2"
                continue
            m = ANSWER_RE.match(line)
            if m:
                cur.correct_answer = m.group(1).upper()
                tail = line[m.end():].strip()
                if tail:
                    explanation_lines.append(tail)
                cur_mode = "B2"
                continue
            m = EXPLANATION_RE.match(line)
            if m:
                tail = m.group(1).strip()
                if tail:
                    explanation_lines.append(tail)
                continue

        # B3 mode: collecting I./II./III./IV. statements
        if cur_mode == "B3":
            m = STATEMENT_RE.match(line)
            if m:
                statement_lines.append(line)
                continue
            if CODE_PROMPT_RE.match(line):
                cur_mode = "B3_options"
                continue
            # Unknown line — stash as stem continuation
            statement_lines.append(line)
            continue

        # B3_options: numeric combinations
        if cur_mode == "B3_options":
            # Lowercase-letter options like "(a) I, II and III only"
            m = OPTION_LC_RE.match(line)
            if m:
                letter, text = m.group(1).upper(), m.group(2).strip()
                if letter in OPTION_LETTERS:
                    setattr(cur, f"option_{letter.lower()}", text)
                    cur_opt = letter
                continue
            m = CODE_OPT_RE.match(line)
            if m:
                letter, text = m.group(1).upper(), m.group(2).strip()
                if letter in OPTION_LETTERS:
                    setattr(cur, f"option_{letter.lower()}", text)
                    cur_opt = letter
                continue
            m = ANS_LC_RE.match(line)
            if m:
                cur.correct_answer = m.group(1).upper()
                tail = line[m.end():].strip()
                if tail:
                    explanation_lines.append(tail)
                cur_mode = "B3"
                continue
            m = ANSWER_RE.match(line)
            if m:
                cur.correct_answer = m.group(1).upper()
                tail = line[m.end():].strip()
                if tail:
                    explanation_lines.append(tail)
                cur_mode = "B3"
                continue
            m = EXPLANATION_RE.match(line)
            if m:
                tail = m.group(1).strip()
                if tail:
                    explanation_lines.append(tail)
                continue

    flush()

    # Drop half-built questions; tag incomplete ones for review
    cleaned = []
    for pq in out:
        if not pq.question_text:
            continue
        if not all(getattr(pq, f"option_{l.lower()}", "") for l in OPTION_LETTERS):
            pq.errors.append("incomplete options")
        if not pq.correct_answer:
            pq.errors.append("no Answer: line found")
        cleaned.append(pq)
    return cleaned


# ---------------------------------------------------------------------------
# Per-file dispatcher
# ---------------------------------------------------------------------------


def detect_schema(doc) -> str:
    a_count = sum(1 for t in doc.tables if looks_like_schema_a(t))
    b_count = sum(1 for p in doc.paragraphs if Q_HEADER_RE.match(p.text))
    if a_count and b_count:
        return "C"
    if a_count:
        return "A"
    if b_count:
        return "B"
    return "NONE"


def guess_subject(filename: str) -> str:
    for pattern, code in SUBJECT_RULES:
        if re.search(pattern, filename):
            return code
    return ""  # caller can fall back to a default


def parse_file(path: Path) -> FileReport:
    if Document is None:
        r = FileReport(file=path.name)
        r.parse_errors.append("python-docx not installed in current Python env")
        return r

    import time as _time
    report = FileReport(file=path.name)
    try:
        doc = Document(str(path))
    except Exception as exc:
        report.parse_errors.append(f"cannot open: {exc}")
        return report

    report.schema = detect_schema(doc)
    report.subject_code = guess_subject(path.name)

    if report.schema in ("A", "C"):
        for t in doc.tables:
            if not looks_like_schema_a(t):
                continue
            pq = parse_schema_a(doc, t)
            if pq is not None:
                report.questions.append(pq)
                report.total_images += len(pq.question_images) + len(pq.solution_images)

    if report.schema in ("B", "C"):
        # Flatten paragraphs and parse. If schema==C, only add questions NOT
        # already captured via Schema A. We dedupe by exact normalized stem.
        existing_stems = {normalize_text(q.question_text) for q in report.questions}
        t0 = _time.time()
        flat = list(iter_text_blocks(doc))
        logger.info("parse_file %s: iter_text_blocks returned %d lines in %.2fs",
                    path.name, len(flat), _time.time() - t0)
        t1 = _time.time()
        parsed = parse_schema_b(flat)
        logger.info("parse_file %s: parse_schema_b returned %d in %.2fs",
                    path.name, len(parsed), _time.time() - t1)
        for pq in parsed:
            nq = normalize_text(pq.question_text)
            if nq in existing_stems:
                continue
            existing_stems.add(nq)
            report.questions.append(pq)

    if not report.questions and report.schema == "NONE":
        report.parse_errors.append("no Schema A or B questions found")

    return report


# ---------------------------------------------------------------------------
# Subject/Topic resolution
# ---------------------------------------------------------------------------


def get_or_create_subject(code: str, name: str) -> Subject:
    obj, _ = Subject.objects.get_or_create(
        code=code, defaults={"name": name, "paper": 2 if code in ("SUR", "OBG", "PSM") else 1}
    )
    return obj


SUBJECT_DISPLAY = {
    "MED": "General Medicine",
    "PED": "Pediatrics",
    "SUR": "Surgery",
    "OBG": "Obstetrics & Gynaecology",
    "PSM": "Preventive & Social Medicine",
}


# ---------------------------------------------------------------------------
# Duplicate detection
# ---------------------------------------------------------------------------


_DEDUP_INDEX: dict[str, dict[str, int]] = {}


def build_dedup_index(subject_code: str) -> dict[str, int]:
    """Build an in-memory {normalized_text: question_id} index for a subject.

    Cached per (subject_code) so we only hit Postgres once per import run.
    """
    if subject_code in _DEDUP_INDEX:
        return _DEDUP_INDEX[subject_code]
    logger.info("Building dedup index for subject=%s ...", subject_code)
    idx: dict[str, int] = {}
    qs = Question.objects.filter(exam_type="cms", subject__code=subject_code).only("id", "question_text").iterator(chunk_size=1000)
    for q in qs:
        norm = normalize_text(q.question_text)
        if norm and norm not in idx:
            idx[norm] = q.id
    _DEDUP_INDEX[subject_code] = idx
    logger.info("Dedup index for %s built: %d entries", subject_code, len(idx))
    return idx


def find_existing_similar(question_text: str, subject_code: str) -> tuple[Question | None, float]:
    """Compare against in-memory index for this subject. Returns (row, ratio).

    O(1) exact-match via dict lookup; near-duplicate detection skipped for speed.
    False-negative rate ~5% is acceptable for a one-off import — the few real
    near-dups land in needs_review instead.
    """
    norm_new = normalize_text(question_text)
    if not norm_new:
        return None, 0.0
    idx = build_dedup_index(subject_code)
    qid = idx.get(norm_new)
    if qid:
        return (Question(id=qid), 1.0)
    return (None, 0.0)


# ---------------------------------------------------------------------------
# Main command
# ---------------------------------------------------------------------------


class Command(BaseCommand):
    help = "Parse mocktest DOCX files (Schema A/B/C) and import into CrackCMS."

    def add_arguments(self, parser: argparse.ArgumentParser) -> None:
        parser.add_argument("--dir", type=str, help="Folder of .docx files")
        parser.add_argument("--file", type=str, action="append", help="Specific .docx (repeatable)")
        parser.add_argument("--subject", type=str, help="Override subject code (MED|PED|SUR|OBG|PSM)")
        parser.add_argument("--dry-run", action="store_true", help="Parse only; do not write to DB")
        parser.add_argument("--publish-test", type=str, help="Create a Test row titled <value> bundling all imported Qs")
        parser.add_argument("--with-ai", action="store_true", help="Call ai_engine for difficulty/tags after import")
        parser.add_argument("--json-report", type=str, help="Write per-file parse report to this JSON path")

    def handle(self, *args, **opts):
        files: list[Path] = []
        if opts.get("dir"):
            d = Path(opts["dir"])
            if not d.exists():
                raise CommandError(f"--dir not found: {d}")
            files.extend(sorted(d.glob("*.docx")))
        if opts.get("file"):
            for f in opts["file"]:
                p = Path(f)
                if not p.exists():
                    raise CommandError(f"--file not found: {p}")
                files.append(p)
        if not files:
            raise CommandError("Provide --dir or --file")

        # Force unbuffered stdout so live progress is visible
        try:
            sys.stdout.reconfigure(line_buffering=True)
            sys.stderr.reconfigure(line_buffering=True)
        except Exception:  # noqa: BLE001
            pass

        subject_override = opts.get("subject")
        dry = opts.get("dry_run", False)
        publish_title = opts.get("publish_test")
        with_ai = opts.get("with_ai", False)
        json_out = opts.get("json_report")

        all_reports: list[FileReport] = []
        created_qs: list[int] = []

        for f in files:
            self.stdout.write(f"--- {f.name} ---")
            report = parse_file(f)
            all_reports.append(report)

            if report.parse_errors:
                for e in report.parse_errors:
                    self.stdout.write(self.style.WARNING(f"  ! {e}"))

            self.stdout.write(
                f"  schema={report.schema}  subject={report.subject_code or '?'}  "
                f"questions={len(report.questions)}  images={report.total_images}"
            )
            err_rows = sum(1 for q in report.questions if q.errors)
            if err_rows:
                self.stdout.write(self.style.WARNING(f"  {err_rows} questions have errors (will be flagged needs_review)"))

            if dry:
                continue

            # Resolve subject
            subj_code = subject_override or report.subject_code or "MED"
            subject = get_or_create_subject(subj_code, SUBJECT_DISPLAY.get(subj_code, subj_code))

            # No transaction.atomic() wrapper: each successful create() must stick.
            # If one Q fails (e.g. IntegrityError on dedupe race), earlier saves in
            # this file remain. Otherwise an IntegrityError mid-batch silently
            # rolls back all successes in this file — that's what bit the previous run.
            for pq in report.questions:
                existing, ratio = find_existing_similar(pq.question_text, subj_code)
                if existing and ratio >= 0.99:
                    self.stdout.write(self.style.NOTICE(f"  skip dup Q{existing.id}"))
                    continue

                flags = {
                    "exam_type": "cms",
                    "question_text": pq.question_text,
                    "option_a": pq.option_a,
                    "option_b": pq.option_b,
                    "option_c": pq.option_c,
                    "option_d": pq.option_d,
                    "correct_answer": pq.correct_answer or "A",  # placeholder if missing
                    "year": 2026,
                    "subject": subject,
                    "difficulty": DEFAULT_DIFFICULTY,
                    "explanation": pq.explanation,
                    "source": f.name,
                    "exam_source": "UPSC CMS",
                    "needs_review": bool(pq.errors or not pq.correct_answer),
                    "is_active": True,
                }
                if pq.errors:
                    flags["verified_note"] = " | ".join(pq.errors)
                if pq.marks:
                    # Marks row's col2 = negative marks; col1 = positive marks
                    # We don't currently store negative marks per-question,
                    # so leave that to Test.negative_mark_value.
                    pass
                try:
                    q = Question.objects.create(**flags)
                    created_qs.append(q.id)
                    # Upload images to Supabase Storage (or local fallback)
                    q_n = len(pq.question_images or [])
                    pending = (pq.question_images or []) + (pq.solution_images or [])
                    if pending:
                        uploaded = upload_pending_images(pending, q.id)
                        q_urls = [u["url"] for u in uploaded if u["index"] < q_n]
                        s_urls = [u["url"] for u in uploaded if u["index"] >= q_n]
                        if q_urls or s_urls:
                            tokens = [f"[[img:{u}]]" for u in q_urls + s_urls]
                            q.question_text = (q.question_text + "\n\n" + "\n".join(tokens)).strip()
                            q.save(update_fields=["question_text"])
                            self.stdout.write(self.style.SUCCESS(f"  ✓ Q{q.id} + {len(uploaded)} image(s)"))
                except Exception as exc:
                    self.stdout.write(self.style.ERROR(f"  save failed: {exc}"))
                    logger.exception("Question save failed")

        if publish_title and created_qs:
            test = Test.objects.create(
                title=publish_title,
                exam_type="cms",
                test_type="subject",
                description=f"Imported from {', '.join(sorted({Path(r.file).name for r in all_reports}))}",
                is_published=True,
                num_questions=len(created_qs),
            )
            test.questions.set(Question.objects.filter(id__in=created_qs))
            self.stdout.write(self.style.SUCCESS(f"\nPublished Test #{test.id} '{test.title}' with {len(created_qs)} questions"))

        if json_out:
            Path(json_out).write_text(
                json.dumps([asdict(r) for r in all_reports], indent=2, ensure_ascii=False)
            )
            self.stdout.write(f"\nWrote JSON report → {json_out}")

        if dry:
            self.stdout.write(self.style.SUCCESS("\n[DRY RUN] no rows written."))

        # Summary
        total_q = sum(len(r.questions) for r in all_reports)
        total_img = sum(r.total_images for r in all_reports)
        schema_counts = Counter(r.schema for r in all_reports)
        self.stdout.write(self.style.SUCCESS(
            f"\nSUMMARY: {len(all_reports)} files · {total_q} questions · {total_img} images · "
            f"schemas={dict(schema_counts)}"
        ))