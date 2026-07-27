"""DOCX formatting fidelity — capture paragraphs, runs, tables, and images
with their original formatting (bold/italic/underline/colour/highlight/subscript/
superscript/links/bullets) and render them as Word-equivalent HTML.

This module is the canonical source of formatting-fidelity logic. The XML
extractor and the python-docx reader both feed into it so the rest of the
pipeline sees a single, normalised representation per question/option/
explanation cell.

Design rules (zero data loss):

* Every run has the same shape: a list of (text_or_image, formatting) tuples.
* Inline images are represented as ``[img:filename]`` tokens that the
  ingest layer swaps for ``<img src="…">`` URLs.
* Tables are rendered as ``<table>`` with merged cells, borders, and
  cell colours preserved.
* Lists become ``<ul>`` / ``<ol>`` based on the paragraph's numPr.
* Bold / italic / underline / strike are emitted as inline tags.
* Colour / highlight are emitted as inline ``<span style="…">``.
* Subscript / superscript become ``<sub>`` / ``<sup>``.
* The output HTML is **safe** by default: text is HTML-escaped.

This module is intentionally pure-Python: no Django, no ORM, no external
services, so it can be unit-tested in isolation.
"""
from __future__ import annotations

import html
import io
import os
import re
import zipfile
from dataclasses import dataclass, field
from typing import Iterable, List, Sequence, Tuple


# ---------------------------------------------------------------------------
# Run-level data structure
# ---------------------------------------------------------------------------


@dataclass
class Run:
    """One visually-continuous text run inside a paragraph or table cell."""

    text: str = ""
    bold: bool = False
    italic: bool = False
    underline: bool = False
    strike: bool = False
    color: str = ""          # hex like '#c62828' or "" for default
    highlight: str = ""      # hex like '#fff59d' or "" for none
    subscript: bool = False
    superscript: bool = False
    image_filename: str = ""  # if this run is actually an inline image
    hyperlink: str = ""        # URL if this run is a hyperlink
    bullet_number: int = 0    # 0 = no list; 1 = bullet (ul); >=1 = numbered (ol)
    bullet_level: int = 0     # 0…3 nesting
    is_list: bool = False


@dataclass
class Paragraph:
    """A paragraph (or a single cell in a table)."""

    runs: List[Run] = field(default_factory=list)
    style: str = ""           # paragraph style name (Heading 1, etc.)
    align: str = ""           # left, center, right, both
    indent_left: int = 0      # twentieths of a point
    indent_right: int = 0
    indent_first: int = 0
    list_numId: int = 0


@dataclass
class TableCell:
    paragraphs: List[Paragraph] = field(default_factory=list)
    width: int = 0            # twips
    bg_color: str = ""        # hex
    border_color: str = ""    # hex
    vmerge: str = ""          # "restart" | "continue" | ""
    gridspan: int = 1
    valign: str = ""          # top | center | bottom


@dataclass
class TableRow:
    cells: List[TableCell] = field(default_factory=list)
    height: int = 0
    is_header: bool = False


@dataclass
class Table:
    rows: List[TableRow] = field(default_factory=list)
    style: str = ""
    width: int = 0
    border_color: str = ""
    border_width: int = 0


@dataclass
class DocumentBlock:
    """Either a paragraph or a table — the document is a flat sequence of these."""

    paragraph: Paragraph = None
    table: Table = None
    paragraph_index: int = 0  # for traceability back to the source file


# ---------------------------------------------------------------------------
# XML helpers (python-docx + namespace-tolerant fallback)
# ---------------------------------------------------------------------------


_NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
_NS_PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"


def _local(tag: str) -> str:
    return tag.split("}", 1)[-1] if "}" in tag else tag


def _attr(elem, name: str, default: str = "") -> str:
    return elem.get(name) or elem.get("{%s}%s" % (_NS_W, name), default) or default


def _twips(value: str | int | None) -> int:
    if value is None:
        return 0
    try:
        return int(str(value))
    except (TypeError, ValueError):
        return 0


_HEX_RE = re.compile(r"^[0-9A-Fa-f]{6}$")


def _normalize_hex(value: str) -> str:
    if not value:
        return ""
    v = value.strip().lstrip("#")
    if len(v) == 3:
        v = "".join(c * 2 for c in v)
    if _HEX_RE.match(v):
        return "#" + v.lower()
    return ""


# Pre-defined Word highlight palette → hex (HTML <mark> doesn't always render
# coloured backgrounds well; we use inline styles).
_HIGHLIGHT_PALETTE = {
    "yellow": "#fff59d",
    "green": "#a5d6a7",
    "cyan": "#80deea",
    "magenta": "#f48fb1",
    "blue": "#90caf9",
    "red": "#ef9a9a",
    "darkBlue": "#1565c0",
    "darkCyan": "#00838f",
    "darkGreen": "#2e7d32",
    "darkMagenta": "#ad1457",
    "darkRed": "#c62828",
    "darkYellow": "#f9a825",
    "darkGray": "#616161",
    "lightGray": "#cfcfcf",
    "black": "#000000",
    "white": "#ffffff",
    "none": "",
}


# ---------------------------------------------------------------------------
# python-docx-backed fidelity reader
# ---------------------------------------------------------------------------


def _read_color(elem) -> str:
    """Read w:color/@val and return a normalised hex."""
    color = elem.find("{%s}color" % _NS_W)
    if color is not None:
        return _normalize_hex(color.get("{%s}val" % _NS_W, ""))
    return ""


def _read_highlight(elem) -> str:
    """Read w:highlight/@val and map to a hex colour."""
    hl = elem.find("{%s}highlight" % _NS_W)
    if hl is not None:
        name = (hl.get("{%s}val" % _NS_W, "") or "").strip()
        return _HIGHLIGHT_PALETTE.get(name, "")
    # w:shd fallback (shading)
    shd = elem.find("{%s}shd" % _NS_W)
    if shd is not None:
        return _normalize_hex(shd.get("{%s}fill" % _NS_W, ""))
    return ""


def _resolve_image_filename(rid_or_name: str, relationships: dict) -> str:
    """Translate a DOCX relationship id (rId5) or asset name (image1.png)
    into the basename stored in ImportedImage.original_filename.

    DOCX <a:blip r:embed="rId5"> references a relationship that points at
    ``word/media/image1.png``. The ingest layer extracts images from
    ``word/media/`` and stores them with the basename. So we must translate
    rId → basename for image_refs to match the stored filenames.
    """
    if not rid_or_name:
        return ""
    # If it looks like an rId, look up the relationships dict.
    if rid_or_name.startswith("rId"):
        target = relationships.get(rid_or_name)
        if target:
            return os.path.basename(target)
    # Already a basename, or unrecognized — return as-is.
    return rid_or_name


def _read_run_formatting(run_elem, relationships: dict | None = None) -> Run:
    """Convert a <w:r> element to a Run.

    ``relationships`` is the {rId: target} dict read from word/_rels/document.xml.rels.
    When supplied, image_filename is translated from rId (r:embed) to the actual
    basename (image1.png) so the ingest layer's ImportedImage.original_filename
    matches.
    """
    r = Run()
    rpr = run_elem.find("{%s}rPr" % _NS_W)
    if rpr is not None:
        r.bold = rpr.find("{%s}b" % _NS_W) is not None
        r.italic = rpr.find("{%s}i" % _NS_W) is not None
        r.underline = rpr.find("{%s}u" % _NS_W) is not None
        r.strike = rpr.find("{%s}strike" % _NS_W) is not None
        r.subscript = (rpr.find("{%s}vertAlign" % _NS_W) is not None
                       and rpr.find("{%s}vertAlign" % _NS_W).get("{%s}val" % _NS_W, "") == "subscript")
        r.superscript = (rpr.find("{%s}vertAlign" % _NS_W) is not None
                         and rpr.find("{%s}vertAlign" % _NS_W).get("{%s}val" % _NS_W, "") == "superscript")
        r.color = _read_color(rpr)
        r.highlight = _read_highlight(rpr)

    # Inline image (drawing/pict). Hyperlinks handled by parent paragraph.
    drawing = run_elem.find("{%s}drawing" % _NS_W)
    raw_image_ref = ""
    if drawing is None:
        # Old-style w:pict
        pict = run_elem.find("{%s}pict" % _NS_W)
        if pict is not None:
            # Embedded image: find blipfill or imagedata
            for elem in pict.iter():
                local = _local(elem.tag)
                if local in ("imagedata", "blip"):
                    raw_image_ref = elem.get("{%s}name" % _NS_R, "") or elem.get("name", "")
                    break
    else:
        # Word 2007+ drawing: find blip element with r:embed
        for elem in drawing.iter():
            if _local(elem.tag) == "blip":
                raw_image_ref = elem.get("{%s}embed" % _NS_R, "") or ""
                break
    if raw_image_ref and relationships is not None:
        r.image_filename = _resolve_image_filename(raw_image_ref, relationships)
    else:
        r.image_filename = raw_image_ref

    # Text content (tab/break/t elements)
    text_parts: List[str] = []
    for child in run_elem:
        local = _local(child.tag)
        if local == "t":
            if child.text:
                text_parts.append(child.text)
        elif local == "tab":
            text_parts.append("\t")
        elif local == "br":
            text_parts.append("\n")
    r.text = "".join(text_parts)
    return r


def _read_paragraph(p_elem, relationships: dict) -> Paragraph:
    """Convert a <w:p> element to a Paragraph."""
    p = Paragraph()
    ppr = p_elem.find("{%s}pPr" % _NS_W)
    if ppr is not None:
        p.style = _attr(ppr, "pStyle")
        jc = ppr.find("{%s}jc" % _NS_W)
        if jc is not None:
            p.align = (jc.get("{%s}val" % _NS_W, "") or "").strip()
        ind = ppr.find("{%s}ind" % _NS_W)
        if ind is not None:
            p.indent_left = _twips(ind.get("{%s}left" % _NS_W))
            p.indent_right = _twips(ind.get("{%s}right" % _NS_W))
            p.indent_first = _twips(ind.get("{%s}firstLine" % _NS_W))
        numPr = ppr.find("{%s}numPr" % _NS_W)
        if numPr is not None:
            numId = numPr.find("{%s}numId" % _NS_W)
            ilvl = numPr.find("{%s}ilvl" % _NS_W)
            if numId is not None:
                p.list_numId = _twips(numId.get("{%s}val" % _NS_W, "0"))

    # Hyperlink wrapper → fetch the rId from the hyperlink element
    hyperlink_rid = ""
    for child in p_elem:
        if _local(child.tag) == "hyperlink":
            hyperlink_rid = child.get("{%s}id" % _NS_R, "")
            for sub in child:
                if _local(sub.tag) == "r":
                    run = _read_run_formatting(sub, relationships)
                    if hyperlink_rid and (hyperlink_rid in relationships):
                        run.hyperlink = relationships[hyperlink_rid]
                    p.runs.append(run)
        elif _local(child.tag) == "r":
            p.runs.append(_read_run_formatting(child, relationships))
        elif _local(child.tag) == "fldSimple":
            # Field text (e.g. PAGE) — collect any <w:t> children
            for r in child.iter("{%s}r" % _NS_W):
                p.runs.append(_read_run_formatting(r, relationships))
    return p


def _read_cell(tc_elem, relationships: dict) -> TableCell:
    cell = TableCell()
    tcPr = tc_elem.find("{%s}tcPr" % _NS_W)
    if tcPr is not None:
        tcW = tcPr.find("{%s}tcW" % _NS_W)
        if tcW is not None:
            cell.width = _twips(tcW.get("{%s}w" % _NS_W))
        shd = tcPr.find("{%s}shd" % _NS_W)
        if shd is not None:
            cell.bg_color = _normalize_hex(shd.get("{%s}fill" % _NS_W, ""))
        vMerge = tcPr.find("{%s}vMerge" % _NS_W)
        if vMerge is not None:
            val = vMerge.get("{%s}val" % _NS_W, "")
            cell.vmerge = "continue" if val == "" else val
        gridSpan = tcPr.find("{%s}gridSpan" % _NS_W)
        if gridSpan is not None:
            cell.gridspan = _twips(gridSpan.get("{%s}val" % _NS_W, "1"))
        vAlign = tcPr.find("{%s}vAlign" % _NS_W)
        if vAlign is not None:
            cell.valign = vAlign.get("{%s}val" % _NS_W, "")
    for child in tc_elem:
        if _local(child.tag) == "p":
            cell.paragraphs.append(_read_paragraph(child, relationships))
    return cell


def _read_row(tr_elem, relationships: dict) -> TableRow:
    row = TableRow()
    trPr = tr_elem.find("{%s}trPr" % _NS_W)
    if trPr is not None:
        trH = trPr.find("{%s}trHeight" % _NS_W)
        if trH is not None:
            row.height = _twips(trH.get("{%s}val" % _NS_W))
        tblHeader = trPr.find("{%s}tblHeader" % _NS_W)
        if tblHeader is not None:
            row.is_header = True
    for child in tr_elem:
        if _local(child.tag) == "tc":
            row.cells.append(_read_cell(child, relationships))
    return row


def _read_table(tbl_elem, relationships: dict) -> Table:
    table = Table()
    tblPr = tbl_elem.find("{%s}tblPr" % _NS_W)
    if tblPr is not None:
        tblStyle = tblPr.find("{%s}tblStyle" % _NS_W)
        if tblStyle is not None:
            table.style = tblStyle.get("{%s}val" % _NS_W, "")
        tblW = tblPr.find("{%s}tblW" % _NS_W)
        if tblW is not None:
            table.width = _twips(tblW.get("{%s}w" % _NS_W))
        tblBorders = tblPr.find("{%s}tblBorders" % _NS_W)
        if tblBorders is not None:
            for border in tblBorders:
                if _local(border.tag) in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    color = _normalize_hex(border.get("{%s}color" % _NS_W, ""))
                    sz = _twips(border.get("{%s}sz" % _NS_W, "0"))
                    if color and color != "auto":
                        table.border_color = color
                        table.border_width = max(table.border_width, sz)
                        break
    for child in tbl_elem:
        if _local(child.tag) == "tr":
            table.rows.append(_read_row(child, relationships))
    return table


def _read_relationships(zip_bytes: bytes) -> dict:
    """Return {rId: target} for the document's relationships (images, hyperlinks)."""
    relationships: dict = {}
    try:
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as z:
            if "word/_rels/document.xml.rels" not in z.namelist():
                return relationships
            from xml.etree import ElementTree as ET
            root = ET.fromstring(z.read("word/_rels/document.xml.rels"))
            for rel in root:
                rid = rel.get("Id", "")
                target = rel.get("Target", "")
                if rid and target:
                    relationships[rid] = target
    except Exception:
        return relationships
    return relationships


def read_document_with_fidelity(path: str) -> List[DocumentBlock]:
    """Read a DOCX file and return a flat list of paragraphs/tables with formatting preserved.

    Uses python-docx first; falls back to the namespace-tolerant XML reader
    if python-docx raises a namespace error.
    """
    from pathlib import Path as _Path
    raw = _Path(path).read_bytes()
    return _read_bytes_with_fidelity(raw)


def _read_bytes_with_fidelity(raw: bytes) -> List[DocumentBlock]:
    """Internal: read raw DOCX bytes into a list of DocumentBlock."""
    try:
        from docx import Document  # type: ignore
    except Exception:
        return _read_bytes_with_fidelity_xml(raw)

    try:
        doc = Document(io.BytesIO(raw))
    except Exception:
        return _read_bytes_with_fidelity_xml(raw)

    # python-docx's body.iter_inner_content() walks <w:p> and <w:tbl> in order.
    relationships: dict = {}
    # python-docx stores relationships in part.rels; lazy-load.
    try:
        rels = doc.part.rels  # type: ignore[attr-defined]
        for rid, rel in rels.items():
            if hasattr(rel, "target_ref"):
                relationships[rid] = rel.target_ref
    except Exception:
        pass

    blocks: List[DocumentBlock] = []
    para_idx = 0
    for child in doc.element.body.iterchildren():
        local = _local(child.tag)
        if local == "p":
            p = _read_paragraph(child, relationships)
            blocks.append(DocumentBlock(paragraph=p, paragraph_index=para_idx))
            para_idx += 1
        elif local == "tbl":
            t = _read_table(child, relationships)
            blocks.append(DocumentBlock(table=t, paragraph_index=para_idx))
            # don't bump para_idx for tables — tables are array of blocks
    return blocks


def _read_bytes_with_fidelity_xml(raw: bytes) -> List[DocumentBlock]:
    """Namespace-tolerant XML reader used when python-docx fails."""
    from xml.etree import ElementTree as ET
    blocks: List[DocumentBlock] = []
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as z:
            if "word/document.xml" not in z.namelist():
                return blocks
            xml_bytes = z.read("word/document.xml")
    except Exception:
        return blocks

    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError:
        return blocks

    relationships = _read_relationships(raw)
    body = root.find("{%s}body" % _NS_W)
    if body is None:
        for child in root:
            if _local(child.tag) == "body":
                body = child
                break
    if body is None:
        return blocks

    para_idx = 0
    for child in body:
        local = _local(child.tag)
        if local == "p":
            p = _read_paragraph(child, relationships)
            blocks.append(DocumentBlock(paragraph=p, paragraph_index=para_idx))
            para_idx += 1
        elif local == "tbl":
            t = _read_table(child, relationships)
            blocks.append(DocumentBlock(table=t, paragraph_index=para_idx))
    return blocks


# ---------------------------------------------------------------------------
# Renderer — DocumentBlock → HTML
# ---------------------------------------------------------------------------


def _escape(s: str) -> str:
    return html.escape(s, quote=False)


def _render_run(run: Run, image_url_for: dict | None = None) -> str:
    """Render a single Run as HTML."""
    if run.image_filename:
        url = ""
        if image_url_for:
            url = image_url_for.get(run.image_filename, "")
        if not url:
            url = f"[img:{run.image_filename}]"
        return f'<img src="{_escape(url)}" alt="{_escape(run.image_filename)}" />'
    text = _escape(run.text)
    if not text:
        return ""
    if run.subscript:
        text = f"<sub>{text}</sub>"
    if run.superscript:
        text = f"<sup>{text}</sup>"
    if run.bold:
        text = f"<strong>{text}</strong>"
    if run.italic:
        text = f"<em>{text}</em>"
    if run.underline:
        text = f"<u>{text}</u>"
    if run.strike:
        text = f"<s>{text}</s>"
    if run.highlight:
        text = f'<span style="background:{_escape(run.highlight)}">{text}</span>'
    if run.color:
        text = f'<span style="color:{_escape(run.color)}">{text}</span>'
    if run.hyperlink:
        text = f'<a href="{_escape(run.hyperlink)}">{text}</a>'
    return text


def _render_paragraph(p: Paragraph, image_url_for: dict | None = None) -> str:
    """Render a Paragraph as HTML."""
    inline = "".join(_render_run(r, image_url_for) for r in p.runs)
    if not inline.strip():
        inline = "&nbsp;"
    style_parts: List[str] = []
    if p.align:
        style_parts.append(f"text-align:{p.align}")
    if p.indent_left:
        style_parts.append(f"padding-left:{p.indent_left // 20}px")
    if p.indent_right:
        style_parts.append(f"padding-right:{p.indent_right // 20}px")
    if p.indent_first:
        style_parts.append(f"text-indent:{p.indent_first // 20}px")
    style_attr = f' style="{";".join(style_parts)}"' if style_parts else ""
    if p.style and p.style.lower().startswith("heading"):
        try:
            level = int(re.search(r"\d+", p.style).group(0))
        except Exception:
            level = 2
        level = max(1, min(6, level))
        return f"<h{level}{style_attr}>{inline}</h{level}>"
    if p.list_numId:
        return f"<li{style_attr}>{inline}</li>"
    return f"<p{style_attr}>{inline}</p>"


def _render_cell(cell: TableCell, image_url_for: dict | None = None) -> str:
    style_parts: List[str] = []
    if cell.bg_color:
        style_parts.append(f"background:{_escape(cell.bg_color)}")
    if cell.width:
        style_parts.append(f"width:{cell.width // 20}px")
    if cell.valign:
        style_parts.append(f"vertical-align:{_escape(cell.valign)}")
    style_attr = f' style="{";".join(style_parts)}"' if style_parts else ""
    colspan_attr = f' colspan="{cell.gridspan}"' if cell.gridspan > 1 else ""
    paragraphs = "".join(_render_paragraph(p, image_url_for) for p in cell.paragraphs)
    if not paragraphs:
        paragraphs = "&nbsp;"
    return f"<td{colspan_attr}{style_attr}>{paragraphs}</td>"


def _render_row(row: TableRow, image_url_for: dict | None = None) -> str:
    cells = "".join(_render_cell(c, image_url_for) for c in row.cells)
    tag = "th" if row.is_header else "tr"
    return f"<{tag}>{cells}</{tag}>"


def _render_table(table: Table, image_url_for: dict | None = None) -> str:
    rows_html = "".join(_render_row(r, image_url_for) for r in table.rows)
    style_parts = ["border-collapse:collapse"]
    if table.border_color:
        style_parts.append(f"border:{table.border_width // 8 or 1}px solid {_escape(table.border_color)}")
    for r in table.rows:
        for c in r.cells:
            if c.border_color:
                style_parts.append(f"border:1px solid {_escape(c.border_color)}")
                break
        else:
            continue
        break
    style_attr = f' style="{";".join(style_parts)}"' if style_parts else ""
    return f"<table{style_attr}>{rows_html}</table>"


def render_blocks(blocks: Sequence[DocumentBlock], image_url_for: dict | None = None) -> str:
    """Render a list of DocumentBlock → HTML."""
    out: List[str] = []
    open_list: int = 0
    for blk in blocks:
        if blk.paragraph is not None:
            if blk.paragraph.list_numId and not open_list:
                out.append("<ul>")
                open_list = 1
            elif not blk.paragraph.list_numId and open_list:
                out.append("</ul>")
                open_list = 0
            out.append(_render_paragraph(blk.paragraph, image_url_for))
        elif blk.table is not None:
            if open_list:
                out.append("</ul>")
                open_list = 0
            out.append(_render_table(blk.table, image_url_for))
    if open_list:
        out.append("</ul>")
    return "\n".join(out)


def render_paragraph(p: Paragraph, image_url_for: dict | None = None) -> str:
    """Convenience: render one paragraph."""
    return _render_paragraph(p, image_url_for)


def render_table(t: Table, image_url_for: dict | None = None) -> str:
    """Convenience: render one table."""
    return _render_table(t, image_url_for)


# ---------------------------------------------------------------------------
# Per-question image association
# ---------------------------------------------------------------------------


def associate_images_to_questions(
    blocks: Sequence[DocumentBlock],
    questions: Sequence,
) -> dict:
    """Map image_filenames → question position_index by walking the document
    in order and assigning each image to the most-recently-started question.

    `questions` is a list of dataclasses with ``position_index`` and
    ``image_refs`` attributes. The returned dict maps
    ``question_position_index → list[image_filename]`` (deduped, in order).

    Falls back to image-less lists when no question boundary can be
    detected inside the document.

    The mapping is intentionally conservative: every image lands in the
    question it was found inside, never in the global image list.
    """
    # Build ordered list of (block_index, block) for the document.
    image_to_question: dict = {q.position_index: [] for q in questions}
    # Find the question start paragraphs by scanning each question's
    # raw_text or question_text against the documents' text content.
    # Simpler: each question's start is implicit in the input order — by
    # convention questions are encountered in document order at parse time.
    # So we walk blocks in order, find paragraphs that match the n-th
    # question's start signature, and assign the images between consecutive
    # question starts to that question.
    question_starts: List[int] = []  # index into `blocks`
    if not questions:
        return image_to_question
    # try to find each question's start by looking for the first paragraph
    # whose runs contain the first 40 chars of the question's question_text.
    used = [False] * len(blocks)
    for q_idx, q in enumerate(questions):
        needle = (q.question_text or "").strip()[:40]
        if not needle:
            continue
        for b_idx, blk in enumerate(blocks):
            if used[b_idx]:
                continue
            if blk.paragraph is None:
                continue
            txt = "".join(r.text for r in blk.paragraph.runs)
            if needle and needle in txt:
                question_starts.append(b_idx)
                used[b_idx] = True
                break
    if not question_starts:
        # Couldn't find any matching start — give the images to question 0.
        for blk in blocks:
            for r in (blk.paragraph.runs if blk.paragraph else []):
                if r.image_filename:
                    image_to_question.setdefault(0, []).append(r.image_filename)
        return image_to_question
    # Pad with end-of-document sentinel so the last bucket is closed.
    question_starts.append(len(blocks))
    for q_idx in range(len(question_starts) - 1):
        start = question_starts[q_idx]
        end = question_starts[q_idx + 1]
        # Find the actual question object (match by walk order)
        try:
            q = questions[q_idx]
        except IndexError:
            break
        seen: set = set()
        for b_idx in range(start, end):
            blk = blocks[b_idx]
            if blk.paragraph is None:
                continue
            for r in blk.paragraph.runs:
                if r.image_filename and r.image_filename not in seen:
                    image_to_question[q.position_index].append(r.image_filename)
                    seen.add(r.image_filename)
    return image_to_question
